"""
Utility module for APIMatic Doc Search
Handles document loading, text processing, embedding generation, and vector store management.
"""

import os
import re
from pathlib import Path
from typing import List, Dict, Tuple, Optional
import logging

# LangChain imports
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS, Chroma
from langchain_core.documents import Document

# Document loading
import PyPDF2
from docx import Document as DocxDocument

from app.config import (
    CHUNK_SIZE,
    CHUNK_OVERLAP,
    SUPPORTED_FORMATS,
    VECTOR_STORE_DIR,
    EMBEDDING_MODELS,
    MARKDOWN_SEPARATORS,
    FAISS_INDEX_PATH,
    EMBEDDING_MODEL_NAME,
)

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────
# Markdown Metadata Extraction
# ─────────────────────────────────────────────
def extract_markdown_metadata(content: str, file_path: Path) -> Dict:
    """
    Extract metadata from a markdown file including frontmatter and heading structure.

    Args:
        content: Raw markdown text
        file_path: Path to the file

    Returns:
        Dictionary of extracted metadata
    """
    metadata = {
        "source": str(file_path),
        "filename": file_path.name,
        "file_type": "markdown",
    }

    # Extract YAML frontmatter if present
    frontmatter_match = re.match(r'^---\s*\n(.*?)\n---\s*\n', content, re.DOTALL)
    if frontmatter_match:
        frontmatter_text = frontmatter_match.group(1)
        for line in frontmatter_text.split('\n'):
            if ':' in line:
                key, _, value = line.partition(':')
                key = key.strip().lower().replace(' ', '_')
                value = value.strip().strip('"').strip("'")
                if key and value:
                    metadata[f"fm_{key}"] = value

    # Extract the first H1 as title if no frontmatter title
    if "fm_title" not in metadata:
        h1_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
        if h1_match:
            metadata["title"] = h1_match.group(1).strip()

    # Extract relative path components for categorization
    try:
        rel_parts = file_path.relative_to(file_path.parents[2]).parts
        if len(rel_parts) > 1:
            metadata["category"] = "/".join(rel_parts[:-1])
    except (ValueError, IndexError):
        pass

    return metadata


def clean_markdown_content(content: str) -> str:
    """
    Clean markdown content for better embedding quality.
    Removes excessive whitespace and formatting artifacts while preserving structure.

    Args:
        content: Raw markdown text

    Returns:
        Cleaned text suitable for embedding
    """
    # Remove frontmatter
    content = re.sub(r'^---\s*\n.*?\n---\s*\n', '', content, flags=re.DOTALL)

    # Remove HTML comments
    content = re.sub(r'<!--.*?-->', '', content, flags=re.DOTALL)

    # Remove image references (keep alt text)
    content = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', r'\1', content)

    # Remove inline HTML tags
    content = re.sub(r'<[^>]+>', '', content)

    # Normalize whitespace (collapse multiple blank lines)
    content = re.sub(r'\n{3,}', '\n\n', content)

    return content.strip()


class DocumentLoader:
    """Handles loading documents from various file formats."""

    @staticmethod
    def load_txt(file_path: Path) -> str:
        """Load text from .txt file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    @staticmethod
    def load_pdf(file_path: Path) -> str:
        """Load text from .pdf file."""
        text = ""
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
        return text

    @staticmethod
    def load_docx(file_path: Path) -> str:
        """Load text from .docx file."""
        try:
            doc = DocxDocument(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            return ""

    @staticmethod
    def load_md(file_path: Path) -> str:
        """Load text from .md file with metadata-aware processing."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error loading MD {file_path}: {e}")
            return ""

    @classmethod
    def load_document(cls, file_path: Path) -> Optional[Document]:
        """
        Load a document from file and return as LangChain Document.

        Args:
            file_path: Path to the document file

        Returns:
            LangChain Document object or None if loading fails
        """
        suffix = file_path.suffix.lower()

        loaders = {
            '.txt': cls.load_txt,
            '.pdf': cls.load_pdf,
            '.docx': cls.load_docx,
            '.md': cls.load_md
        }

        if suffix not in loaders:
            logger.warning(f"Unsupported format: {suffix}")
            return None

        raw_text = loaders[suffix](file_path)
        if not raw_text.strip():
            return None

        # Enhanced metadata for markdown files
        if suffix == '.md':
            metadata = extract_markdown_metadata(raw_text, file_path)
            text = clean_markdown_content(raw_text)
        else:
            metadata = {"source": str(file_path), "filename": file_path.name}
            text = raw_text

        if text.strip():
            return Document(page_content=text, metadata=metadata)
        return None

    @classmethod
    def load_documents_from_directory(cls, directory: Path) -> Tuple[List[Document], Dict]:
        """
        Load all supported documents from a directory (recursively).

        Args:
            directory: Path to directory containing documents

        Returns:
            Tuple of (list of Documents, statistics dictionary)
        """
        documents = []
        stats = {
            'total_files': 0,
            'loaded_files': 0,
            'failed_files': 0,
            'total_size_bytes': 0,
            'file_types': {}
        }

        # Skip hidden directories and common non-doc dirs
        skip_dirs = {'.git', 'node_modules', '__pycache__', '.venv', 'venv'}

        for file_path in Path(directory).rglob('*'):
            # Skip files in hidden/excluded directories
            if any(part in skip_dirs for part in file_path.parts):
                continue

            if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_FORMATS:
                stats['total_files'] += 1
                stats['total_size_bytes'] += file_path.stat().st_size

                # Track file type
                ext = file_path.suffix.lower()
                stats['file_types'][ext] = stats['file_types'].get(ext, 0) + 1

                doc = cls.load_document(file_path)
                if doc:
                    documents.append(doc)
                    stats['loaded_files'] += 1
                else:
                    stats['failed_files'] += 1

        logger.info(f"Loaded {stats['loaded_files']}/{stats['total_files']} documents")
        return documents, stats


class TextProcessor:
    """Handles text chunking and processing with markdown-aware splitting."""

    def __init__(self, chunk_size: int = CHUNK_SIZE, chunk_overlap: int = CHUNK_OVERLAP):
        """
        Initialize text processor with markdown-aware separators.

        Args:
            chunk_size: Size of text chunks in characters
            chunk_overlap: Overlap between chunks in characters
        """
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=MARKDOWN_SEPARATORS
        )

    def split_documents(self, documents: List[Document]) -> List[Document]:
        """
        Split documents into smaller chunks, preserving metadata.

        Args:
            documents: List of LangChain Documents

        Returns:
            List of chunked Documents
        """
        chunks = self.text_splitter.split_documents(documents)

        # Enrich chunk metadata with position info
        doc_chunk_counts = {}
        for chunk in chunks:
            source = chunk.metadata.get("source", "unknown")
            doc_chunk_counts[source] = doc_chunk_counts.get(source, 0) + 1
            chunk.metadata["chunk_index"] = doc_chunk_counts[source]

        # Second pass — add total chunk counts
        for chunk in chunks:
            source = chunk.metadata.get("source", "unknown")
            chunk.metadata["total_chunks"] = doc_chunk_counts[source]

        logger.info(f"Split {len(documents)} documents into {len(chunks)} chunks")
        return chunks


class EmbeddingEngine:
    """Handles embedding generation using Hugging Face models."""

    _instance = None
    _current_model = None

    @classmethod
    def get_instance(cls, model_key: str = None) -> 'EmbeddingEngine':
        """
        Get or create a singleton embedding engine.
        Avoids reloading the model on every request.

        Args:
            model_key: Key from EMBEDDING_MODELS config

        Returns:
            EmbeddingEngine singleton instance
        """
        if model_key is None:
            model_key = "all-mpnet-base-v2"

        if cls._instance is None or cls._current_model != model_key:
            cls._instance = cls(model_key)
            cls._current_model = model_key
        return cls._instance

    def __init__(self, model_key: str):
        """
        Initialize embedding engine with specified model.

        Args:
            model_key: Key from EMBEDDING_MODELS config
        """
        if model_key not in EMBEDDING_MODELS:
            raise ValueError(f"Unknown model: {model_key}")

        self.model_key = model_key
        self.model_name = EMBEDDING_MODELS[model_key]["name"]

        logger.info(f"Loading embedding model: {self.model_name}")
        self.embeddings = HuggingFaceEmbeddings(
            model_name=self.model_name,
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """Embed a list of documents."""
        return self.embeddings.embed_documents(texts)

    def embed_query(self, text: str) -> List[float]:
        """Embed a single query."""
        return self.embeddings.embed_query(text)


class VectorStoreManager:
    """Manages vector store creation, saving, and loading."""

    def __init__(self, store_type: str, embedding_engine: EmbeddingEngine):
        """
        Initialize vector store manager.

        Args:
            store_type: Type of vector store ('FAISS' or 'ChromaDB')
            embedding_engine: Initialized EmbeddingEngine instance
        """
        self.store_type = store_type
        self.embedding_engine = embedding_engine
        self.vector_store = None

    def create_vector_store(self, documents: List[Document]) -> None:
        """
        Create vector store from documents.

        Args:
            documents: List of LangChain Document chunks
        """
        logger.info(f"Creating {self.store_type} vector store from {len(documents)} chunks")

        if self.store_type == "FAISS":
            self.vector_store = FAISS.from_documents(
                documents=documents,
                embedding=self.embedding_engine.embeddings
            )
        elif self.store_type == "ChromaDB":
            self.vector_store = Chroma.from_documents(
                documents=documents,
                embedding=self.embedding_engine.embeddings,
                persist_directory=str(VECTOR_STORE_DIR / "chroma_db")
            )
        else:
            raise ValueError(f"Unsupported vector store: {self.store_type}")

        logger.info(f"{self.store_type} vector store created successfully")

    def save_vector_store(self, name: str) -> None:
        """
        Save vector store to disk.

        Args:
            name: Name for the saved store
        """
        if self.vector_store is None:
            raise ValueError("No vector store to save")

        save_path = VECTOR_STORE_DIR / f"{name}_{self.store_type.lower()}"

        if self.store_type == "FAISS":
            self.vector_store.save_local(str(save_path))
            logger.info(f"FAISS store saved to {save_path}")
        elif self.store_type == "ChromaDB":
            # ChromaDB persists automatically if persist_directory is set
            logger.info(f"ChromaDB persisted to {VECTOR_STORE_DIR / 'chroma_db'}")

    def load_vector_store(self, name: str) -> bool:
        """
        Load vector store from disk.

        Args:
            name: Name of the saved store

        Returns:
            True if loaded successfully, False otherwise
        """
        load_path = VECTOR_STORE_DIR / f"{name}_{self.store_type.lower()}"

        try:
            if self.store_type == "FAISS":
                self.vector_store = FAISS.load_local(
                    str(load_path),
                    self.embedding_engine.embeddings,
                    allow_dangerous_deserialization=True
                )
                logger.info(f"FAISS store loaded from {load_path}")
                return True
            elif self.store_type == "ChromaDB":
                self.vector_store = Chroma(
                    persist_directory=str(VECTOR_STORE_DIR / "chroma_db"),
                    embedding_function=self.embedding_engine.embeddings
                )
                logger.info(f"ChromaDB loaded from {VECTOR_STORE_DIR / 'chroma_db'}")
                return True
        except Exception as e:
            logger.error(f"Error loading vector store: {e}")
            return False

        return False

    def similarity_search(self, query: str, k: int = 5) -> List[Tuple[Document, float]]:
        """
        Perform similarity search on vector store.

        Args:
            query: Search query
            k: Number of top results to return

        Returns:
            List of (Document, similarity_score) tuples
        """
        if self.vector_store is None:
            raise ValueError("No vector store loaded")

        # Perform search with scores
        results = self.vector_store.similarity_search_with_score(query, k=k)

        logger.info(f"Found {len(results)} results for query: {query[:50]}...")
        return results


# Convenience function for quick setup
def create_semantic_search_system(
    data_directory: Path,
    embedding_model: str,
    vector_store_type: str
) -> Tuple[VectorStoreManager, Dict]:
    """
    Create complete semantic search system from directory.

    Args:
        data_directory: Path to directory with documents
        embedding_model: Embedding model key
        vector_store_type: Type of vector store

    Returns:
        Tuple of (VectorStoreManager, statistics)
    """
    # Load documents
    documents, stats = DocumentLoader.load_documents_from_directory(data_directory)

    if not documents:
        raise ValueError("No documents loaded from directory")

    # Process documents
    processor = TextProcessor()
    chunks = processor.split_documents(documents)

    # Create embeddings
    embedding_engine = EmbeddingEngine(embedding_model)

    # Create and populate vector store
    vector_manager = VectorStoreManager(vector_store_type, embedding_engine)
    vector_manager.create_vector_store(chunks)

    stats['total_chunks'] = len(chunks)
    return vector_manager, stats
